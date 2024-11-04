import csv
import docx
import subprocess
import os
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import tkinter as tk
from tkinter import ttk, messagebox

class OfficeScriptGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("The Office Script Generator")
        self.root.geometry("400x300")
        
        # Create main frame with padding
        main_frame = ttk.Frame(root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Season selection
        ttk.Label(main_frame, text="Season:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.season_var = tk.StringVar()
        season_combo = ttk.Combobox(main_frame, textvariable=self.season_var, width=5)
        season_combo['values'] = tuple(range(1, 10))  # Seasons 1-9
        season_combo.grid(row=0, column=1, sticky=tk.W, pady=5)
        
        # Episodes input
        ttk.Label(main_frame, text="Episodes (comma-separated):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.episodes_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.episodes_var, width=15).grid(row=1, column=1, sticky=tk.W, pady=5)
        
        # Generate button
        ttk.Button(main_frame, text="Generate Scripts", command=self.generate_scripts).grid(row=2, column=0, columnspan=2, pady=20)
        
        # Status label
        self.status_var = tk.StringVar()
        ttk.Label(main_frame, textvariable=self.status_var, wraplength=350).grid(row=3, column=0, columnspan=2, pady=5)

    def generate_scripts(self):
        season = self.season_var.get()
        episodes = self.episodes_var.get().strip()
        
        if not season or not episodes:
            messagebox.showerror("Error", "Please enter both season and episode numbers")
            return
            
        try:
            episodes_list = episodes.split(',')
            # Validate input
            season_num = int(season)
            episodes_nums = [int(ep.strip()) for ep in episodes_list]
            
            self.status_var.set("Generating scripts...")
            self.root.update()
            
            self.create_document(str(season_num), episodes_list)
            
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numbers")
            return

    def create_document(self, season, episodes):
        try:
            # Create a new Word document
            doc = docx.Document()

            # Set margins (0.5 inches = 457200 EMUs)
            sections = doc.sections
            for section in sections:
                section.left_margin = 457200
                section.right_margin = 457200
                section.top_margin = 457200
                section.bottom_margin = 457200

            # Read the CSV file
            with open('The-Office-Lines-V4.csv', 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                
                # Filter and format lines
                current_episode = None
                for row in reader:
                    if row['season'] == season and row['episode'] in episodes:
                        # Add episode header if it's a new episode
                        if current_episode != row['episode']:
                            current_episode = row['episode']
                            heading = doc.add_heading(f"Season {season}, Episode {current_episode}: {row['title']}", level=1)
                            doc.add_paragraph()  # Add some spacing
                        
                        # Format the line with bold speaker
                        paragraph = doc.add_paragraph()
                        speaker_part = paragraph.add_run(f"{row['speaker']}: ")
                        speaker_part.bold = True
                        paragraph.add_run(row['line'])
                        
                        # Minimize spacing between lines
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0)
                        paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

            # Create filename based on season and episodes
            filename_base = f"office_s{season.zfill(2)}e{'-'.join(episodes)}"
            downloads_path = os.path.expanduser("~/Downloads")
            docx_filename = os.path.join(downloads_path, f"{filename_base}.docx")

            # Save the Word document
            doc.save(docx_filename)
            
            self.status_var.set(f"Success! File saved as:\n{docx_filename}")
            
        except Exception as e:
            self.status_var.set(f"Error: {str(e)}")
            messagebox.showerror("Error", str(e))

def main():
    root = tk.Tk()
    app = OfficeScriptGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()