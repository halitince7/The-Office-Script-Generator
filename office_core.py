import csv
import docx
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import os

def create_script_document(season, episodes):
    """
    Core function to create Office script document.
    Returns (success, message, filename)
    """
    try:
        # Create a new Word document
        doc = docx.Document()

        # Set margins (0.5 inches = 457200 EMUs)
        sections = doc.sections
        for section in sections:
            section.left_margin = 457200
            section.right_margin = 457200
            section.top_margin = 457200
            section.bottom_margin = 457200

        # Read the CSV file
        episodes_found = set()
        current_episode = None
        current_scene = None
        scene_counter = 1  # Initialize scene counter
        
        with open('The-Office-Lines-V4.csv', 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            
            for row in reader:
                if row['season'] == season and row['episode'] in episodes:
                    episodes_found.add(row['episode'])
                    # Add episode header if it's a new episode
                    if current_episode != row['episode']:
                        current_episode = row['episode']
                        heading = doc.add_heading(f"Season {season}, Episode {current_episode}: {row['title']}", level=1)
                        doc.add_paragraph()  # Add some spacing
                        current_scene = None  # Reset scene tracking
                        scene_counter = 1  # Reset scene counter for new episode
                    
                    # Add scene header if it's a new scene
                    if current_scene != row['scene']:
                        current_scene = row['scene']
                        scene_heading = doc.add_heading(f"Scene {scene_counter}", level=2)
                        scene_heading.style.font.size = Pt(12)  # Make scene headers smaller
                        scene_counter += 1  # Increment scene counter
                    
                    # Format the line with bold speaker
                    paragraph = doc.add_paragraph()
                    speaker_part = paragraph.add_run(f"{row['speaker']}: ")
                    speaker_part.bold = True
                    paragraph.add_run(row['line'])
                    
                    # Minimize spacing between lines
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

        # Check if all episodes were found
        missing_episodes = set(episodes) - episodes_found
        if missing_episodes:
            return False, f"Episodes not found: {', '.join(missing_episodes)}", None

        # Create filename and save
        filename_base = f"office_s{season.zfill(2)}e{'-'.join(episodes)}"
        downloads_path = os.path.expanduser("~/Downloads")
        docx_filename = os.path.join(downloads_path, f"{filename_base}.docx")
        
        doc.save(docx_filename)
        return True, f"Successfully saved to: {docx_filename}", docx_filename

    except Exception as e:
        return False, f"Error: {str(e)}", None