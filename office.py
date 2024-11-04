import csv
import docx
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import sys
import os

# Check command line arguments
if len(sys.argv) != 3:
    print("Usage: python script.py season episode_numbers")
    print("Example: python script.py 5 13,14,15")
    sys.exit(1)

season = sys.argv[1]
episodes = sys.argv[2].split(',')

# Create a new Word document
doc = docx.Document()

# Set margins (0.5 inches = 457200 EMUs)
sections = doc.sections
for section in sections:
    section.left_margin = 457200
    section.right_margin = 457200
    section.top_margin = 457200
    section.bottom_margin = 457200

# Read the CSV file
with open('The-Office-Lines-V4.csv', 'r', encoding='utf-8') as file:
    reader = csv.DictReader(file)
    
    # Filter and format lines
    current_episode = None
    episodes_found = set()  # Track which episodes we find
    
    for row in reader:
        if row['season'] == season and row['episode'] in episodes:
            episodes_found.add(row['episode'])  # Add to our found episodes
            # Add episode header if it's a new episode
            if current_episode != row['episode']:
                current_episode = row['episode']
                heading = doc.add_heading(f"Season {season}, Episode {current_episode}: {row['title']}", level=1)
                doc.add_paragraph()  # Add some spacing
            
            # Format the line with bold speaker
            paragraph = doc.add_paragraph()
            speaker_part = paragraph.add_run(f"{row['speaker']}: ")
            speaker_part.bold = True
            paragraph.add_run(row['line'])
            
            # Minimize spacing between lines
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    # After processing all rows, check if we found all requested episodes
    missing_episodes = set(episodes) - episodes_found
    if missing_episodes:
        print(f"Error: The following episodes were not found in the CSV: {', '.join(missing_episodes)}")
        sys.exit(1)

# Create filename based on season and episodes
filename_base = f"office_s{season.zfill(2)}e{'-'.join(episodes)}"
downloads_path = os.path.expanduser("~/Downloads")  # Get Downloads folder path
docx_filename = os.path.join(downloads_path, f"{filename_base}.docx")

# Save the Word document to Downloads folder
try:
    doc.save(docx_filename)
    print(f"Successfully saved to: {docx_filename}")
except Exception as e:
    print(f"Error saving file: {str(e)}")


